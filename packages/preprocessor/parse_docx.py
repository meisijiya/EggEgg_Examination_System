"""DOCX 资料解析器(支持任意科目 — 公司战略等)。

把 DOCX 文档拆解为"段落/小节"记录,每条带 source_file + needs_ai_answer=true,
供 downstream multi-agent pipeline 生成结构化题目。

设计原则:
- **零 LLM 边界**:仅使用 python-docx + mammoth + 正则启发式
- **段落粒度**:用 Heading 风格 + 空行做分节
- **失败显式**:解析失败写入 errors.log,绝不静默吞错
- **向后兼容**:CLI 不传参数 = 走财务原 pipeline(无 DOCX 时空退出)
"""
from __future__ import annotations

import argparse
import json
import logging
import re
import sys
from pathlib import Path
from typing import Any

from pydantic import BaseModel, ConfigDict, Field

# ---------------------------------------------------------------------------
# 路径常量(默认 = 项目根 / 公司战略和风险管理,但 CLI 可覆盖)
# ---------------------------------------------------------------------------

PROJECT_ROOT = Path(__file__).resolve().parents[2]
DOCX_DIR_DEFAULT = PROJECT_ROOT / "公司战略和风险管理"
PARSED_DIR = PROJECT_ROOT / "data" / "parsed"
ERRORS_LOG = PARSED_DIR / "errors.log"

# ---------------------------------------------------------------------------
# 日志
# ---------------------------------------------------------------------------

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger("parse_docx")

_ERRORS_FILE_HANDLE: Any = None


def _log_error(docx_name: str, msg: str) -> None:
    """DOCX 解析错误到 errors.log(累计所有调用)。"""
    global _ERRORS_FILE_HANDLE
    if _ERRORS_FILE_HANDLE is None:
        ERRORS_LOG.parent.mkdir(parents=True, exist_ok=True)
        _ERRORS_FILE_HANDLE = open(ERRORS_LOG, "a", encoding="utf-8")
    _ERRORS_FILE_HANDLE.write(f"[{docx_name}]: {msg}\n")
    _ERRORS_FILE_HANDLE.flush()


# ---------------------------------------------------------------------------
# Pydantic Schema — 段落记录
# ---------------------------------------------------------------------------


class DocSegment(BaseModel):
    """DOCX 一个段落/小节的标准化记录(extra='forbid' 严格校验)。

    字段对齐 deepwork Phase 1.2 任务清单:
    - source_file:原 DOCX 文件相对路径(可追溯)
    - paragraph_index:在源文件中的段落序号(0-indexed)
    - heading_style:段落标题样式(Heading 1/2/Title 等;非标题段落=None)
    - raw_text:段落原始文本(纯文本,无图)
    - section_path:层级路径(由 Heading 标题拼接,例如 '战略选择 > 案例 > 1')
    - needs_ai_answer:固定 True(DOCX 资料型题目无标准答案,需 AI 生成)
    """

    model_config = ConfigDict(extra="forbid")

    id: str = Field(..., min_length=8, description="稳定 UUID(PARAGRAPH_HASH)")
    source_file: str = Field(..., min_length=1, description="源 DOCX 文件名")
    paragraph_index: int = Field(..., ge=0, description="段落序号")
    heading_style: str | None = Field(default=None, description="段落标题样式")
    raw_text: str = Field(..., min_length=1, description="段落纯文本")
    section_path: str = Field(default="", description="标题层级路径")
    needs_ai_answer: bool = Field(default=True, description="是否需要 AI 补答案")


# ---------------------------------------------------------------------------
# DOCX 段落提取(基于 python-docx)
# ---------------------------------------------------------------------------


def _iter_paragraphs(docx_path: Path) -> list[tuple[int, str | None, str]]:
    """返回 [(paragraph_index, heading_style_or_None, text), ...]。

    - paragraph_index:0-indexed 段落序号
    - heading_style:段落 style.name 含 'Heading'/'Title' 时为该样式;否则 None
    - text:段落纯文本(strip 后的非空字符串)
    """
    from docx import Document  # python-docx

    results: list[tuple[int, str | None, str]] = []
    doc = Document(str(docx_path))
    for idx, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        if not text:
            continue
        style_name = para.style.name if para.style is not None else ""
        heading_style: str | None = None
        if "Heading" in style_name or "Title" in style_name:
            heading_style = style_name
        results.append((idx, heading_style, text))
    return results


def _iter_table_texts(docx_path: Path) -> list[tuple[int, str]]:
    """提取 DOCX 中的表格内容,作为段落补充。返回 [(paragraph_index, text), ...]。

    索引从 100_000 起,避免与正段落冲突(同一 DOCX 段落数远小于 100k)。
    """
    from docx import Document

    results: list[tuple[int, str]] = []
    doc = Document(str(docx_path))
    base_idx = 100_000
    for t_idx, table in enumerate(doc.tables):
        rows_text: list[str] = []
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            rows_text.append(" | ".join(cells))
        if rows_text:
            results.append((base_idx + t_idx, "\n".join(rows_text)))
    return results


def _build_section_paths(
    paragraphs: list[tuple[int, str | None, str]],
) -> list[str]:
    """为每个段落计算层级 section_path(由 Heading 堆叠)。

    规则:
    - Heading 1 段落:section_path = h1_text
    - Heading 2:section_path = h1 > h2
    - 普通段落:section_path = 最近上级 Heading 路径
    """
    paths: list[str] = []
    h_stack: list[str] = []  # 当前 Heading 栈

    for _, heading_style, text in paragraphs:
        if heading_style is None or not heading_style:
            # 沿用最近的 heading 栈
            paths.append(" > ".join(h_stack))
            continue
        # Heading 1 / Title 推入栈顶
        if "Heading 1" in heading_style or heading_style == "Title":
            h_stack = [text]
        elif "Heading 2" in heading_style:
            h_stack = h_stack[:1] + [text]
        elif "Heading 3" in heading_style:
            h_stack = h_stack[:2] + [text]
        else:
            h_stack = h_stack[:3] + [text]
        paths.append(" > ".join(h_stack))
    return paths


def _compute_id(source_file: str, paragraph_index: int, text: str) -> str:
    """生成稳定 UUID(基于 source + index + text 前 100 字 hash)。"""
    import hashlib

    h = hashlib.sha256(f"{source_file}|{paragraph_index}|{text[:100]}".encode("utf-8")).hexdigest()
    return h[:16]


def parse_one_docx(docx_path: Path) -> list[DocSegment]:
    """解析单个 DOCX → 段落实例列表。

    跳过空段落(已 strip),只保留非空纯文本 + 表格内容。
    """
    source_file = docx_path.name
    out: list[DocSegment] = []

    paragraphs = _iter_paragraphs(docx_path)
    table_segments = _iter_table_texts(docx_path)

    if not paragraphs and not table_segments:
        _log_error(source_file, "无任何段落或表格内容")
        return []

    section_paths = _build_section_paths(paragraphs)
    for (idx, heading_style, text), section_path in zip(paragraphs, section_paths):
        out.append(
            DocSegment(
                id=_compute_id(source_file, idx, text),
                source_file=source_file,
                paragraph_index=idx,
                heading_style=heading_style,
                raw_text=text,
                section_path=section_path,
                needs_ai_answer=True,
            )
        )

    # 表格作为独立段落追加(paragraph_index=100_000+t_idx)
    for t_idx, t_text in table_segments:
        out.append(
            DocSegment(
                id=_compute_id(source_file, t_idx, t_text),
                source_file=source_file,
                paragraph_index=t_idx,
                heading_style=None,
                raw_text=f"[表格]\n{t_text}",
                section_path="",
                needs_ai_answer=True,
            )
        )

    return out


def write_jsonl(items: list[DocSegment], path: Path) -> int:
    """写 JSONL(extras-fail via DocSegment.model_dump);返回写入行数。"""
    path.parent.mkdir(parents=True, exist_ok=True)
    n = 0
    with open(path, "w", encoding="utf-8") as f:
        for seg in items:
            f.write(seg.model_dump_json(ensure_ascii=False) + "\n")
            n += 1
    return n


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def _build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        prog="parse_docx",
        description="DOCX → segments.jsonl 解析器(支持任意科目)",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument(
        "--docx-dir",
        type=Path,
        default=DOCX_DIR_DEFAULT,
        help="DOCX 输入目录",
    )
    parser.add_argument(
        "--subject",
        type=str,
        default="corporate_strategy",
        help="科目代码,影响输出文件命名(default=corporate_strategy → data/parsed/corporate_strategy_questions_docx.jsonl)",
    )
    parser.add_argument(
        "--output-jsonl",
        type=Path,
        default=None,
        help="合并输出 JSONL 路径(None = 按 subject 自动决定)",
    )
    parser.add_argument(
        "--glob",
        type=str,
        default="*.docx",
        help="DOCX 文件匹配 glob(默认 *.docx)",
    )
    return parser


def _safe_rel(path: Path, base: Path) -> Path:
    try:
        return path.relative_to(base)
    except ValueError:
        return path


def main(argv: list[str] | None = None) -> int:
    """CLI 入口。

    使用方式:
        # 公司战略 6 DOCX
        python -m packages.preprocessor.parse_docx
        # 任意科目
        python -m packages.preprocessor.parse_docx \\
            --docx-dir 'path/to/docx' \\
            --subject my-subject
    """
    parser = _build_arg_parser()
    args = parser.parse_args(argv)

    docx_dir: Path = args.docx_dir
    subject: str = args.subject

    if args.output_jsonl is not None:
        merged_path: Path = args.output_jsonl
    else:
        merged_path = PARSED_DIR / f"{subject}_questions_docx.jsonl"

    if not docx_dir.exists():
        logger.error(f"DOCX 目录不存在: {docx_dir}")
        return 1

    docx_paths = sorted(docx_dir.glob(args.glob))
    if not docx_paths:
        logger.error(f"DOCX 目录无 {args.glob} 文件: {docx_dir}")
        return 1

    merged_path.parent.mkdir(parents=True, exist_ok=True)

    logger.info(
        "parse_docx: subject=%s docx_dir=%s → %s (n_files=%d)",
        subject,
        _safe_rel(docx_dir, PROJECT_ROOT),
        _safe_rel(merged_path, PROJECT_ROOT),
        len(docx_paths),
    )

    all_segments: list[DocSegment] = []
    failed: list[tuple[str, str]] = []

    for docx_path in docx_paths:
        logger.info(f"开始解析 {docx_path.name}")
        try:
            segs = parse_one_docx(docx_path)
        except Exception as e:  # noqa: BLE001
            logger.error(f"DOCX {docx_path.name} 解析异常: {e}")
            _log_error(docx_path.name, f"异常: {e!r}")
            failed.append((docx_path.name, str(e)))
            continue
        if not segs:
            failed.append((docx_path.name, "无段落"))
            continue
        all_segments.extend(segs)
        logger.info(f"  {docx_path.name}: {len(segs)} 段")

    if all_segments:
        n = write_jsonl(all_segments, merged_path)
        logger.info(f"合并输出: {n} 段 → {_safe_rel(merged_path, PROJECT_ROOT)}")

    # 收尾
    global _ERRORS_FILE_HANDLE
    if _ERRORS_FILE_HANDLE is not None:
        _ERRORS_FILE_HANDLE.close()
        _ERRORS_FILE_HANDLE = None

    print("\n" + "=" * 60)
    print("DOCX 解析完成报告")
    print("=" * 60)
    print(f"输入 DOCX: {len(docx_paths)}")
    print(f"输出段落: {len(all_segments)} 段")
    print(f"失败 DOCX: {len(failed)}")
    for name, reason in failed:
        print(f"  - {name}: {reason}")
    print(f"输出 JSONL: {_safe_rel(merged_path, PROJECT_ROOT)}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
